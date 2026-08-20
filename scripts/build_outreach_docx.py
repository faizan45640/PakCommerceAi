import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def create_outreach_docx():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Base Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(6)

    # Document Header / Title
    title = doc.add_paragraph()
    title_run = title.add_run("PakCommerce AI — Seller Outreach & Data Guide")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F) # Navy
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Step-by-step field guide for connecting with Pakistani online stores and collecting delivery history")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    subtitle.paragraph_format.space_after = Pt(14)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(10)
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="CCCCCC"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(p_div_border)

    # -------------------------------------------------------------
    # 1. Payout & Compensation
    # -------------------------------------------------------------
    h1 = doc.add_paragraph()
    r = h1.add_run("1. Your Payout & Earnings Structure")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(6)

    doc.add_paragraph(
        "You earn a direct reward for every verified store dataset (Excel or CSV file) you collect and upload:"
    )

    # Payout Table
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Store Data Size (Order / Customer History)", "Payout Per Store"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = Inches(3.2)
        set_cell_background(cell, "1E3A5F")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data = [
        ("Under 10,000 orders / customers in past history", "Rs. 3,000"),
        ("10,000+ orders / customers in past history", "Rs. 5,000")
    ]

    for row_idx, (size_text, payout_text) in enumerate(data, start=1):
        c0 = table.cell(row_idx, 0)
        c1 = table.cell(row_idx, 1)
        c0.width = Inches(3.2)
        c1.width = Inches(3.2)
        set_cell_background(c0, "F8F9FA" if row_idx % 2 == 1 else "FFFFFF")
        set_cell_background(c1, "EBF3FC" if row_idx % 2 == 1 else "F0F7FF")
        set_cell_margins(c0, top=140, bottom=140, left=180, right=180)
        set_cell_margins(c1, top=140, bottom=140, left=180, right=180)

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p0.add_run(size_text)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(payout_text)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x0F, 0x68, 0x36) # Green

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(6)
    p_note.paragraph_format.space_after = Pt(12)
    r_note = p_note.add_run("Payment is cleared directly once the file is uploaded and verified.")
    r_note.font.size = Pt(10)
    r_note.font.italic = True
    r_note.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # -------------------------------------------------------------
    # 2. What Are We Doing & Why Sellers Share Data
    # -------------------------------------------------------------
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. What Are We Doing? (In Simple Words)")
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        "Every online seller in Pakistan deals with Cash on Delivery (COD). Their biggest pain point is that 20% to 30% of buyers reject parcels at their doorstep, causing sellers to lose thousands of rupees in return delivery fees every month."
    )
    doc.add_paragraph(
        "We are building PakCommerce AI (a university project from UET Lahore) to help sellers predict high-risk returns and choose the best courier (Trax, PostEx, Leopards, TCS) for every city."
    )

    doc.add_paragraph(
        "Why will sellers share past order data with you?\n"
        "We offer them a 100% Free Delivery & Courier Audit Report. In return for their past order file, our team creates a custom visual report for them showing:\n"
        "  • Which cities have their worst return and cancellation rates.\n"
        "  • Which courier company is causing them the biggest losses.\n"
        "  • Practical recommendations to save shipping costs."
    )

    # -------------------------------------------------------------
    # 3. Where to Find 25–30 Stores Daily
    # -------------------------------------------------------------
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. Where to Find 25–30 Stores Daily")
    r3.font.size = Pt(14)
    r3.font.bold = True
    r3.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        "1. Instagram Hashtags:\n"
        "   Search hashtags like #clothingbrandpk, #onlineshoppingpakistan, #modestfashionpk, #shoesstorepk, #unstitchedcollection.\n"
        "   Look for active pages with 2,000 to 50,000 followers.\n\n"
        "2. Suggested Accounts on Instagram:\n"
        "   When you open any Pakistani brand page, tap the small down-arrow next to 'Follow' — Instagram will automatically suggest 15–20 similar Pakistani brand pages.\n\n"
        "3. Look for their WhatsApp Link:\n"
        "   Most brand bios have a direct WhatsApp Business link (wa.me link). You can message them on WhatsApp or send an Instagram DM."
    )

    # -------------------------------------------------------------
    # 4. What Data We Need vs. What We DO NOT Want
    # -------------------------------------------------------------
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. Data Privacy: What We Need vs. What We DO NOT Want")
    r4.font.size = Pt(14)
    r4.font.bold = True
    r4.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        "Always reassure sellers: \"We only analyze city delivery rates. We do NOT want your customers' private details.\""
    )

    box_table = doc.add_table(rows=1, cols=2)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_table.autofit = False

    c_yes = box_table.cell(0, 0)
    c_no = box_table.cell(0, 1)
    c_yes.width = Inches(3.2)
    c_no.width = Inches(3.2)

    set_cell_background(c_yes, "F0FDF4") # light green
    set_cell_background(c_no, "FEF2F2") # light red
    set_cell_margins(c_yes, top=140, bottom=140, left=180, right=180)
    set_cell_margins(c_no, top=140, bottom=140, left=180, right=180)

    p_yes = c_yes.paragraphs[0]
    p_yes.paragraph_format.space_after = Pt(0)
    r_yes_title = p_yes.add_run("✓ What We Need (Standard columns):\n")
    r_yes_title.font.bold = True
    r_yes_title.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
    p_yes.add_run(
        "• Order Date\n"
        "• Destination City (Lahore, KHI, etc.)\n"
        "• Order Total (Rs.)\n"
        "• Item Category (Clothes, Shoes, etc.)\n"
        "• Courier Name (Trax, PostEx, TCS, etc.)\n"
        "• Delivery Status (Delivered / Returned)"
    )

    p_no = c_no.paragraphs[0]
    p_no.paragraph_format.space_after = Pt(0)
    r_no_title = p_no.add_run("✗ What We DO NOT Want:\n")
    r_no_title.font.bold = True
    r_no_title.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
    p_no.add_run(
        "• Customer Full Names\n"
        "• Complete House / Street Addresses\n"
        "• Bank account / Credit card details\n"
        "• Profit margins or product costs\n\n"
        "(Sellers can delete these columns!)"
    )

    # -------------------------------------------------------------
    # 5. Message Templates (Copy & Paste)
    # -------------------------------------------------------------
    h5 = doc.add_paragraph()
    r5 = h5.add_run("5. Ready-to-Use Message Templates")
    r5.font.size = Pt(14)
    r5.font.bold = True
    r5.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(4)

    # Template 1: WhatsApp / Instagram DM (Roman Urdu)
    doc.add_paragraph("Template 1: WhatsApp / Instagram DM (Roman Urdu) — Best for local brand owners:").runs[0].font.bold = True

    t1_box = doc.add_table(rows=1, cols=1)
    t1_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1 = t1_box.cell(0, 0)
    c1.width = Inches(6.5)
    set_cell_background(c1, "F8FAFC")
    set_cell_margins(c1, top=160, bottom=160, left=200, right=200)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    p1.add_run(
        "\"AOA! Hope you're doing well. 😊\n\n"
        "I came across your store page, really liked your collection!\n\n"
        "Actually, we are building PakCommerce AI (a tech project from UET Lahore) to help Pakistani online sellers reduce COD parcel returns and courier shipping losses.\n\n"
        "We are doing Free Delivery & Courier Audits for local brands this week. If you can share a past order export file (Excel/CSV from your Trax, PostEx, or Shopify portal — strictly no customer names or payment info needed), our team will analyze it and send you a free report showing:\n"
        "  1. Which cities have your worst return & cancellation rates.\n"
        "  2. Which courier is costing you the most return delivery losses.\n"
        "  3. Simple recommendations to save shipping costs.\n\n"
        "Would you like us to run a free check for your store? Let me know and I can send you the quick 2-minute export steps!\""
    )

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(4)
    p_space.paragraph_format.space_after = Pt(0)

    # Template 2: Professional English
    doc.add_paragraph("Template 2: Professional English (LinkedIn / Email):").runs[0].font.bold = True

    t2_box = doc.add_table(rows=1, cols=1)
    t2_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    c2 = t2_box.cell(0, 0)
    c2.width = Inches(6.5)
    set_cell_background(c2, "F8FAFC")
    set_cell_margins(c2, top=160, bottom=160, left=200, right=200)
    p2 = c2.paragraphs[0]
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(
        "\"Hi [Name / Brand],\n\n"
        "Hope you're doing well.\n\n"
        "I'm reaching out from PakCommerce AI, a project built at UET Lahore focusing on helping Pakistani ecommerce sellers reduce Cash on Delivery (COD) return rates and courier losses.\n\n"
        "We are currently running free Delivery Performance & Courier Audits for selected local brands. By looking at an export of your past shipments (no customer personal details or payment info required), we create a custom report showing your return hotspots and courier success rates by city.\n\n"
        "Would you be open to receiving a free delivery audit report for your store?\n\n"
        "Best regards,\n"
        "PakCommerce AI Team\""
    )

    # -------------------------------------------------------------
    # 6. Quick Export Instructions for the Seller
    # -------------------------------------------------------------
    h6 = doc.add_paragraph()
    r6 = h6.add_run("6. When the Seller Says \"Yes, how do I export?\"")
    r6.font.size = Pt(14)
    r6.font.bold = True
    r6.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h6.paragraph_format.space_before = Pt(12)
    h6.paragraph_format.space_after = Pt(4)

    doc.add_paragraph("Copy and paste these simple steps to them:")

    t3_box = doc.add_table(rows=1, cols=1)
    t3_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    c3 = t3_box.cell(0, 0)
    c3.width = Inches(6.5)
    set_cell_background(c3, "F8FAFC")
    set_cell_margins(c3, top=160, bottom=160, left=200, right=200)
    p3 = c3.paragraphs[0]
    p3.paragraph_format.space_after = Pt(0)
    p3.add_run(
        "\"Great! You can export your order file in 2 minutes:\n\n"
        "• From Courier Portal (Trax, PostEx, Leopards, TCS):\n"
        "  Log in to your portal ➔ Go to Orders / Consignment Reports ➔ Select past 6 to 12 months ➔ Click 'Export as Excel or CSV'.\n\n"
        "• From Shopify:\n"
        "  Go to Shopify Admin ➔ Orders ➔ Click 'Export' on top right ➔ Select 'All Orders' or pick a date range ➔ Export as CSV.\n\n"
        "• From WooCommerce:\n"
        "  Go to WooCommerce ➔ Orders or Analytics ➔ Export / Download CSV.\n\n"
        "(Note: You are welcome to delete customer names or private address details if you prefer — we only need City, Courier, and Delivered/Returned status).\n\n"
        "You can drop the file right here on WhatsApp!\""
    )

    # -------------------------------------------------------------
    # 7. Handling Common Questions & Objections
    # -------------------------------------------------------------
    h7 = doc.add_paragraph()
    r7 = h7.add_run("7. Cheat Sheet: Answering Questions & Objections")
    r7.font.size = Pt(14)
    r7.font.bold = True
    r7.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h7.paragraph_format.space_before = Pt(12)
    h7.paragraph_format.space_after = Pt(4)

    faqs = [
        ("Q: \"Is my customer data safe?\"",
         "\"100% safe! We do not need customer names, street addresses, or payment details. We only look at city, order amount, courier name, and delivered or returned status. You can remove any private columns before sending.\""),
        ("Q: \"Are you guys competitors or planning to sell products?\"",
         "\"No, we are university students and software developers building operational tools for sellers. We do not sell any physical products or run retail stores.\""),
        ("Q: \"How long will the report take?\"",
         "\"Our team delivers the visual audit report back to you within 24 to 48 hours.\""),
        ("Q: \"Is this really free?\"",
         "\"Yes, completely free. All we ask in return is your honest feedback on the report.\"" )
    ]

    for q, a in faqs:
        p_faq = doc.add_paragraph()
        p_faq.paragraph_format.space_before = Pt(3)
        p_faq.paragraph_format.space_after = Pt(3)
        r_q = p_faq.add_run(q + "\n")
        r_q.font.bold = True
        r_q.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        r_a = p_faq.add_run("Answer: " + a)
        r_a.font.italic = True

    # -------------------------------------------------------------
    # 8. What to Do When You Receive a File
    # -------------------------------------------------------------
    h8 = doc.add_paragraph()
    r8 = h8.add_run("8. What to Do When You Receive a File")
    r8.font.size = Pt(14)
    r8.font.bold = True
    r8.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    h8.paragraph_format.space_before = Pt(12)
    h8.paragraph_format.space_after = Pt(4)

    doc.add_paragraph(
        "1. Rename the file clearly:\n"
        "   Format: [BrandName]_[CourierOrPlatform]_[Date].xlsx\n"
        "   (Example: ModestVibe_Trax_Aug2026.xlsx)\n\n"
        "2. Note it in the tracking sheet:\n"
        "   • Brand Name & Instagram/Website link\n"
        "   • Contact Person WhatsApp number\n"
        "   • Number of orders in the file (e.g. 3,500 orders or 12,000 orders)\n"
        "   • Date received\n\n"
        "3. Upload the file to the shared Google Drive folder.\n\n"
        "4. Your payout will be calculated based on the order count (Rs. 3,000 for <10k, Rs. 5,000 for 10k+) and cleared upon verification."
    )

    # Output file
    output_path = Path("docs/outreach/Seller_Outreach_Guide.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Successfully generated: {output_path.resolve()}")

if __name__ == "__main__":
    create_outreach_docx()
